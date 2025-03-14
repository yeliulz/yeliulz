from flask import Flask, request, jsonify, render_template, send_file
import os
import logging
from pdf2docx import Converter
from docx import Document
import pandas as pd
from dotenv import load_dotenv
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from datetime import datetime, timedelta
import threading
import time
import json

load_dotenv()

app = Flask(__name__)

# Set up logging configuration
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_file():
    full_name = request.form.get('fullName')
    email = request.form.get('email')
    file = request.files.get('file')

    if not full_name or not email or not file:
        return jsonify({"error": "Full name, email, and file are required."}), 400

    file_path = os.path.join('uploads', file.filename)
    file.save(file_path)

    try:
        if file.filename.endswith('.pdf'):
            docx_file = convert_pdf_to_docx(file_path)
        elif file.filename.endswith('.xlsx'):
            docx_file = convert_excel_to_docx(file_path)
        else:
            return jsonify({"error": "Invalid file type. Only PDF and Excel files are allowed."}), 400

        # Forward data to Google Sheets
        forward_to_google_sheets(full_name, email, file.filename)

        # Send the converted file directly for download
        return send_file(docx_file, as_attachment=True)

    except Exception as e:
        logging.error(f"Error during file conversion: {str(e)}")
        return jsonify({"error": f"Failed to convert file: {str(e)}"}), 500

def convert_pdf_to_docx(pdf_path):
    docx_path = pdf_path.replace('.pdf', '.docx')
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)  # Convert all pages
        cv.close()
        logging.info(f"Successfully converted PDF to DOCX: {docx_path}")
        return docx_path
    except Exception as e:
        logging.error(f"Failed to convert PDF: {str(e)}")
        raise

def convert_excel_to_docx(excel_path):
    docx_path = excel_path.replace('.xlsx', '.docx')
    try:
        df = pd.read_excel(excel_path, sheet_name=None)
        doc = Document()
        for sheet_name, data in df.items():
            doc.add_heading(sheet_name, level=1)
            table = doc.add_table(rows=1, cols=len(data.columns))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, column_name in enumerate(data.columns):
                hdr_cells[i].text = column_name
            for index, row in data.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
        doc.save(docx_path)
        logging.info(f"Successfully converted Excel to DOCX: {docx_path}")
        return docx_path
    except Exception as e:
        logging.error(f"Failed to convert Excel: {str(e)}")
        raise

def get_google_credentials():
    try:
        google_credentials_json = os.getenv('GOOGLE_CREDENTIALS_JSON')
        if not google_credentials_json:
            raise ValueError("GOOGLE_CREDENTIALS_JSON environment variable not set")
        
        credentials_info = json.loads(google_credentials_json)
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds = ServiceAccountCredentials.from_json_keyfile_dict(credentials_info, scope)
        return creds
    except Exception as e:
        logging.error(f"Failed to load Google credentials: {str(e)}")
        raise

def forward_to_google_sheets(full_name, email, file_name):
    try:
        creds = get_google_credentials()
        client = gspread.authorize(creds)

        # Open the Google Sheet
        sheet = client.open("file converter").sheet1

        # Append the data with a timestamp
        timestamp = datetime.now().isoformat()
        sheet.append_row([full_name, email, file_name, timestamp])
        logging.info("Data successfully forwarded to Google Sheets.")
    except Exception as e:
        logging.error(f"Failed to add data to Google Sheets: {str(e)}")

def summarize_conversions():
    try:
        creds = get_google_credentials()
        client = gspread.authorize(creds)

        # Open the Google Sheet
        sheet = client.open("file converter").sheet1

        # Get all records
        records = sheet.get_all_records()

        # Filter records from the last 24 hours
        now = datetime.now()
        last_24_hours = [record for record in records if datetime.fromisoformat(record['timestamp']) > now - timedelta(hours=24)]

        # Count the conversions
        conversion_count = len(last_24_hours)

        # Write the summary to Google Sheets
        summary_sheet = client.open("file converter").worksheet("Summary")
        summary_sheet.append_row([now.isoformat(), conversion_count])
        logging.info("Summary successfully written to Google Sheets.")
    except Exception as e:
        logging.error(f"Failed to summarize conversions: {str(e)}")

# Schedule the summary function to run every 24 hours
def schedule_summary():
    while True:
        summarize_conversions()
        time.sleep(86400)  # Sleep for 24 hours

if __name__ == '__main__':
    # Start the summary scheduler in a separate thread
    threading.Thread(target=schedule_summary).start()
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
